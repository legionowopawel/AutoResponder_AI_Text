"""
responders/analiza.py
Responder KEYWORDS3 — analiza powtórzeń tekstu.
Przyjmuje DOCX (base64) lub czysty tekst.
Zwraca przetworzony DOCX z kolorowymi podświetleniami powtarzających się słów.
Brak odpowiedzi AI — tylko mechaniczna analiza.
"""
import os
import io
import re
import base64
from collections import defaultdict
from flask import current_app

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import RGBColor
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

# ── Paleta kolorów (identyczna jak w 1.py) ────────────────────────────────────
_COLOR_DATA = [
    (WD_COLOR_INDEX.YELLOW,      (255, 255,   0), False),
    (WD_COLOR_INDEX.BRIGHT_GREEN,(  0, 255,   0), False),
    (WD_COLOR_INDEX.TURQUOISE,   (  0, 255, 255), False),
    (WD_COLOR_INDEX.PINK,        (255,   0, 255), False),
    (WD_COLOR_INDEX.BLUE,        (  0,   0, 255), True),
    (WD_COLOR_INDEX.RED,         (255,   0,   0), True),
    (WD_COLOR_INDEX.DARK_YELLOW, (128, 128,   0), True),
    (WD_COLOR_INDEX.TEAL,        (  0, 128, 128), True),
    (WD_COLOR_INDEX.VIOLET,      (128,   0, 128), True),
    (WD_COLOR_INDEX.GREEN,       (  0, 128,   0), True),
    (WD_COLOR_INDEX.DARK_BLUE,   (  0,   0, 128), True),
    (WD_COLOR_INDEX.DARK_RED,    (128,   0,   0), True),
    (WD_COLOR_INDEX.GRAY_50,     (128, 128, 128), True),
    (WD_COLOR_INDEX.GRAY_25,     (192, 192, 192), False),
]
_PALETTE = (_COLOR_DATA * 11)[:150]

X_VAL = 2200  # domyślny zasięg powtórzeń w znakach


def _get_smart_root(word: str) -> str:
    """Uproszczona stemizacja dla języka polskiego."""
    w = word.lower()
    if len(w) < 3:
        return w
    if "śmie" in w:
        if "śmier" in w: return "śmierć"
        if "śmiet" in w: return "śmietana"
        return "śmiech/uśmiech"
    suffixes = [
        'ego', 'emu', 'ach', 'ami', 'ych', 'ich', 'owi', 'om',
        'em', 'am', 'ia', 'ie', 'y', 'a', 'u', 'i',
        'ł', 'ła', 'li', 'ły', 'cie', 'sz'
    ]
    stem = w
    for s in sorted(suffixes, key=len, reverse=True):
        if stem.endswith(s) and len(stem) - len(s) >= 3:
            stem = stem[:-len(s)]
            break
    return stem


def _build_highlight_map(full_text: str, x_val: int = X_VAL) -> dict:
    """
    Buduje słownik {pozycja_znaku: kolor} dla powtarzających się słów.
    """
    matches = list(re.finditer(r'\b\w+\b', full_text))
    groups  = defaultdict(list)
    for m in matches:
        w = m.group().lower()
        if len(w) > 2:
            groups[_get_smart_root(w)].append(m)

    highlight_map = {}
    color_idx = 0

    for root, m_list in groups.items():
        if len(m_list) < 2:
            continue
        group_has_rep = False
        for i, m1 in enumerate(m_list):
            for j, m2 in enumerate(m_list):
                if i != j and abs(m1.start() - m2.start()) <= x_val:
                    group_has_rep = True
                    if m2.start() not in highlight_map:
                        highlight_map[m2.start()] = _PALETTE[color_idx % len(_PALETTE)]
        if group_has_rep:
            if m_list[0].start() not in highlight_map:
                highlight_map[m_list[0].start()] = _PALETTE[color_idx % len(_PALETTE)]
            color_idx += 1

    return highlight_map


def _apply_highlights_to_doc(doc: "Document", highlight_map: dict) -> "Document":
    """
    Nakłada podświetlenia na akapity dokumentu.
    """
    full_text = "".join([p.text + " \n " for p in doc.paragraphs])

    for para in doc.paragraphs:
        txt = para.text
        if not txt:
            continue
        # Znajdź pozycję akapitu w pełnym tekście
        try:
            para_start = full_text.index(txt)
        except ValueError:
            continue

        tokens = re.split(r'(\b\w+\b)', txt)
        para.clear()
        offset = 0

        for t in tokens:
            run = para.add_run(t)
            pos = para_start + offset
            if pos in highlight_map:
                color_info = highlight_map[pos]
                run.font.highlight_color = color_info[0]
                if color_info[2]:  # ciemny kolor → biały tekst
                    run.font.color.rgb = RGBColor(255, 255, 255)
            offset += len(t)

    return doc


def _text_to_doc(text: str) -> "Document":
    """Tworzy nowy dokument DOCX z czystego tekstu."""
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    return doc


def _doc_to_base64(doc: "Document") -> str:
    """Zapisuje dokument do pamięci i zwraca base64."""
    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode("ascii")


def build_analiza_section(body: str, attachment_b64: str = None,
                           attachment_name: str = None) -> dict:
    """
    Buduje sekcję 'analiza':
    - Jeśli attachment_b64 podany → analizuje przesłany DOCX
    - Jeśli brak załącznika → tworzy DOCX z treści maila i analizuje
    Zwraca DOCX z podświetleniami jako base64.
    Brak odpowiedzi AI.
    """
    if not _HAS_DOCX:
        current_app.logger.error("python-docx nie jest zainstalowane!")
        return {
            "reply_html": "<p>Błąd serwera: brak biblioteki python-docx.</p>",
            "docx": None,
        }

    try:
        # ── Wczytaj lub utwórz dokument ───────────────────────────────────────
        if attachment_b64:
            try:
                docx_bytes = base64.b64decode(attachment_b64)
                doc = Document(io.BytesIO(docx_bytes))
                source = "załącznik"
            except Exception as e:
                current_app.logger.warning("Błąd wczytania DOCX: %s — analizuję treść maila", e)
                doc    = _text_to_doc(body)
                source = "treść maila (błąd DOCX)"
        else:
            doc    = _text_to_doc(body)
            source = "treść maila"

        # ── Buduj mapę podświetleń ────────────────────────────────────────────
        full_text     = "".join([p.text + " \n " for p in doc.paragraphs])
        highlight_map = _build_highlight_map(full_text, X_VAL)

        # ── Nałóż podświetlenia ───────────────────────────────────────────────
        doc = _apply_highlights_to_doc(doc, highlight_map)

        # ── Zapisz do base64 ──────────────────────────────────────────────────
        docx_b64 = _doc_to_base64(doc)

        current_app.logger.info(
            "Analiza: źródło=%s | powtórzenia=%d grup",
            source, len(set(str(v) for v in highlight_map.values()))
        )

        return {
            "reply_html": (
                "<p style='color:#555; font-size:10px;'>"
                f"Analiza powtórzeń wykonana na podstawie: {source}.<br>"
                "Wynik w załączniku DOCX."
                "</p>"
            ),
            "docx": {
                "base64":   docx_b64,
                "filename": "analiza_powtorzen.docx",
            },
        }

    except Exception as e:
        current_app.logger.exception("Błąd analizy: %s", e)
        return {
            "reply_html": "<p>Wystąpił błąd podczas analizy powtórzeń.</p>",
            "docx": None,
        }
